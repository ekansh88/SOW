from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def replace_rich_text_in_placeholder(doc, placeholder, new_text):
    """
    Replace the rich text placeholder in the doc with the new rich text content.
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Clear the runs inside the paragraph with the placeholder
            paragraph.clear()

            # Create a new run for the new rich text content
            run = paragraph.add_run()
            
            # Example of adding rich text (bold and italic):
            run.bold = True  # Set the text bold
            run.italic = True  # Set the text italic
            run.font.size = Pt(12)  # Set font size

            # Add new text
            run.text = new_text

            break

# Load the document
doc = Document(r'C:\Codes\ROPs\SOW\Sow\RichTextPy\templateTesting.docx')

# Replace the rich text in the placeholder
replace_rich_text_in_placeholder(doc, '{{Project Background}}', 'This is the new rich text content.')

# Save the updated document
doc.save(r'C:\Codes\ROPs\SOW\Sow\RichTextPy\templateTestingOutput.docx')
